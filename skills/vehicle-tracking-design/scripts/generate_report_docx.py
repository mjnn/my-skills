#!/usr/bin/env python3
"""Convert UE validation markdown report to styled docx using python-docx."""

from __future__ import annotations

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor
except ImportError:
    print("ERROR: python-docx required. Run: python -m pip install python-docx", file=sys.stderr)
    raise


HEADER_FILL = "2F5496"
ZEBRA_FILL = "F2F2F2"


def _set_cell_shading(cell, hex_color: str) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        from docx.oxml import OxmlElement

        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)


def _add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        padded = row + [""] * (ncols - len(row))
        for c_idx, text in enumerate(padded):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = text
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10.5)
                    run.font.name = "微软雅黑"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            if r_idx == 0:
                _set_cell_shading(cell, HEADER_FILL)
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
            elif r_idx % 2 == 0:
                _set_cell_shading(cell, ZEBRA_FILL)


def _parse_table(lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in lines:
        if re.match(r"^\|[-:\s|]+\|$", line.strip()):
            continue
        rows.append([c.strip() for c in line.strip().strip("|").split("|")])
    return rows


def markdown_to_docx(md: str, output: Path) -> None:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "微软雅黑"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    lines = md.replace("\r\n", "\n").split("\n")
    i = 0
    while i < len(lines):
        stripped = lines[i].strip()
        if not stripped:
            i += 1
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=0)
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=1)
            i += 1
            continue
        if stripped.startswith("|"):
            tbl_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                tbl_lines.append(lines[i].strip())
                i += 1
            _add_table(doc, _parse_table(tbl_lines))
            continue
        if stripped.startswith("- "):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
            i += 1
            continue
        m = re.match(r"\*\*(.+?)\*\*:\s*(.*)", stripped)
        if m:
            p = doc.add_paragraph()
            r1 = p.add_run(m.group(1) + ": ")
            r1.bold = True
            p.add_run(m.group(2))
            i += 1
            continue
        if "**" in stripped:
            p = doc.add_paragraph()
            parts = re.split(r"(\*\*.+?\*\*)", stripped)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    p.add_run(part[2:-2]).bold = True
                elif part:
                    p.add_run(part)
            i += 1
            continue
        doc.add_paragraph(stripped)
        i += 1

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))


def main() -> int:
    if len(sys.argv) != 3:
        print(
            "Usage: python generate_report_docx.py <input.md> <output.docx>",
            file=sys.stderr,
        )
        return 1
    md_path = Path(sys.argv[1])
    output = Path(sys.argv[2])
    if not md_path.is_file():
        print(f"Input not found: {md_path}", file=sys.stderr)
        return 1
    markdown_to_docx(md_path.read_text(encoding="utf-8"), output)
    print(f"Generated: {output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
